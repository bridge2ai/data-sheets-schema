#!/usr/bin/env python3
"""
Unit tests for preprocess_sources.py

Tests preprocessing functionality (PDF/HTML extraction).
"""

import unittest
import tempfile
import shutil
from pathlib import Path
import sys

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

try:
    from src.download.preprocess_sources import (
        extract_pdf_text,
        extract_html_text,
        preprocess_project,
        preprocess_manifest,
    )
except ImportError as e:
    print(f"Warning: Could not import preprocess_sources: {e}", file=sys.stderr)
    extract_pdf_text = None
    extract_html_text = None
    preprocess_project = None
    preprocess_manifest = None


class TestPreprocessing(unittest.TestCase):
    """Test preprocessing functionality."""

    def setUp(self):
        """Set up test fixtures."""
        if extract_html_text is None:
            self.skipTest("preprocess_sources module not available")

        self.test_dir = tempfile.mkdtemp()
        self.test_path = Path(self.test_dir)
        self.src_dir = self.test_path / "src"
        self.dst_dir = self.test_path / "dst"
        self.src_dir.mkdir()
        self.dst_dir.mkdir()

    def tearDown(self):
        """Clean up temporary directory."""
        if hasattr(self, 'test_dir'):
            shutil.rmtree(self.test_dir)

    def test_extract_html_basic(self):
        """Test basic HTML text extraction."""
        html_content = """
        <html>
            <head><title>Test</title></head>
            <body>
                <h1>Header</h1>
                <p>This is a test paragraph.</p>
            </body>
        </html>
        """
        html_file = self.src_dir / "test.html"
        html_file.write_text(html_content)

        extracted_text = extract_html_text(html_file)

        self.assertIn("Header", extracted_text)
        self.assertIn("test paragraph", extracted_text)
        self.assertNotIn("<html>", extracted_text)  # Tags should be removed

    def test_extract_html_removes_scripts(self):
        """Test that script and style tags are removed."""
        html_content = """
        <html>
            <head>
                <style>body { color: red; }</style>
                <script>alert('test');</script>
            </head>
            <body>
                <p>Visible content</p>
                <script>console.log('hidden');</script>
            </body>
        </html>
        """
        html_file = self.src_dir / "test_scripts.html"
        html_file.write_text(html_content)

        extracted_text = extract_html_text(html_file)

        self.assertIn("Visible content", extracted_text)
        self.assertNotIn("color: red", extracted_text)
        self.assertNotIn("alert", extracted_text)
        self.assertNotIn("console.log", extracted_text)

    def test_extract_html_error_handling(self):
        """Test HTML extraction error handling."""
        nonexistent_file = self.src_dir / "nonexistent.html"

        # Should return empty string on error, not crash
        result = extract_html_text(nonexistent_file)
        self.assertEqual(result, "")

    def test_preprocess_project_txt_files(self):
        """Test preprocessing of .txt files (should be copied as-is)."""
        # Create test .txt file
        txt_file = self.src_dir / "test.txt"
        txt_content = "This is plain text content.\n"
        txt_file.write_text(txt_content)

        stats = preprocess_project(self.src_dir, self.dst_dir)

        # Should have copied the file
        self.assertEqual(stats['copied'], 1)

        # Output file should exist and match content
        output_file = self.dst_dir / "test.txt"
        self.assertTrue(output_file.exists())
        self.assertEqual(output_file.read_text(), txt_content)

    def test_preprocess_project_json_files(self):
        """Test preprocessing of .json files (should be preserved)."""
        json_file = self.src_dir / "test.json"
        json_content = '{"key": "value"}'
        json_file.write_text(json_content)

        stats = preprocess_project(self.src_dir, self.dst_dir)

        self.assertEqual(stats['copied'], 1)

        output_file = self.dst_dir / "test.json"
        self.assertTrue(output_file.exists())
        self.assertEqual(output_file.read_text(), json_content)

    def test_preprocess_project_html_files(self):
        """Test preprocessing of .html files (should extract text)."""
        html_file = self.src_dir / "test.html"
        html_content = """<html><body>
            <h1>Test HTML Document</h1>
            <p>This is a test paragraph with sufficient content to pass the minimum character threshold.</p>
            <p>The HTML extraction requires more than 100 characters of text content.</p>
            <p>This third paragraph ensures we have enough content for the test to pass successfully.</p>
        </body></html>"""
        html_file.write_text(html_content)

        stats = preprocess_project(self.src_dir, self.dst_dir)

        # Should have extracted HTML
        self.assertEqual(stats['html_extracted'], 1)

        # Output should be .txt file
        output_file = self.dst_dir / "test.txt"
        self.assertTrue(output_file.exists())

        output_content = output_file.read_text()
        self.assertIn("Test HTML Document", output_content)
        self.assertIn("sufficient content", output_content)
        self.assertGreater(len(output_content), 100)  # Verify meets minimum threshold

    def test_preprocess_project_skips_directories(self):
        """Test that subdirectories are skipped."""
        subdir = self.src_dir / "subdir"
        subdir.mkdir()
        (subdir / "nested.txt").write_text("Nested content")

        preprocess_project(self.src_dir, self.dst_dir)

        # Should not process files in subdirectories
        self.assertFalse((self.dst_dir / "subdir" / "nested.txt").exists())

    def test_preprocess_project_nonexistent_directory(self):
        """Test handling of nonexistent source directory."""
        nonexistent = self.test_path / "nonexistent"

        stats = preprocess_project(nonexistent, self.dst_dir)

        # Should return stats with no operations
        self.assertEqual(stats['copied'], 0)
        self.assertEqual(stats['html_extracted'], 0)

    def test_preprocess_creates_output_directory(self):
        """Test that output directory is created if it doesn't exist."""
        new_dst = self.test_path / "new_destination"

        # Create a source file
        (self.src_dir / "test.txt").write_text("content")

        # Preprocess to non-existent destination
        stats = preprocess_project(self.src_dir, new_dst)

        # Destination should have been created
        self.assertTrue(new_dst.exists())
        self.assertEqual(stats['copied'], 1)

    def test_preprocess_manifest_builds_text_only_outputs(self):
        raw_root = self.test_path / "raw"
        project_raw = raw_root / "AI_READI"
        project_raw.mkdir(parents=True)
        (project_raw / "source.md").write_text(
            "# Source  \r\n \tIndented content  \r\n" + "content  \r\n" * 100
        )
        manifest_file = self.test_path / "manifest.yaml"
        manifest_file.write_text(
            """
version: 1
default_minimum_characters: 100
projects:
  AI_READI:
    - id: documentation
      source_type: documentation
      url: https://example.org/docs
      raw_file: source.md
      processed_file: documentation.txt
      curation_note: Verified source note
      verification_url: https://example.org/verification
""".strip()
        )
        output_root = self.test_path / "processed"

        stats = preprocess_manifest(
            manifest_file,
            raw_root,
            output_root,
            ["AI_READI"],
        )

        self.assertEqual(stats["errors"], 0)
        output = output_root / "AI_READI" / "documentation.txt"
        self.assertTrue(output.exists())
        output_text = output.read_text()
        self.assertIn("Source URL: https://example.org/docs", output_text)
        # #421: the curation note must NOT reach the bundle. It is manifest
        # metadata addressed to a curator, and writing it here put
        # conflict-resolution guidance ("prefer this over X where the two
        # disagree") and, for AI-READI, a DOI that appears in no source
        # document in front of the model. It stays in source_manifest.yaml.
        self.assertNotIn("Curation note:", output_text)
        self.assertNotIn("Verified source note", output_text)
        # Nor the verification URL (#427). #421 removed the note and left this
        # behind, so curator text was still reaching the model — 7 lines across
        # four of the five bundles, three of them AI_READI's and none CHORUS's.
        # It records where a curator checked a capture against upstream and
        # appears in no source document, so a record could cite it and still be
        # scored grounded. It stays in source_manifest.yaml.
        self.assertNotIn("Verification URL:", output_text)
        self.assertNotIn(" \n", output_text)
        self.assertNotIn(" \t", output_text)
        self.assertEqual(
            [path.suffix for path in output.parent.iterdir()],
            [".txt"],
        )

    def test_preprocess_manifest_rejects_short_extraction(self):
        raw_root = self.test_path / "raw"
        project_raw = raw_root / "VOICE"
        project_raw.mkdir(parents=True)
        (project_raw / "stub.txt").write_text("stub")
        manifest_file = self.test_path / "manifest.yaml"
        manifest_file.write_text(
            """
version: 1
default_minimum_characters: 100
projects:
  VOICE:
    - id: documentation
      raw_file: stub.txt
      processed_file: documentation.txt
""".strip()
        )

        stats = preprocess_manifest(
            manifest_file,
            raw_root,
            self.test_path / "processed",
            ["VOICE"],
        )

        self.assertEqual(stats["errors"], 1)


class DocxStructuredDocumentTags(unittest.TestCase):
    """#886: runs, cells and rows wrapped in `w:sdt` (a Google Docs export
    artefact) were dropped by python-docx's `.text`; the bundle read "tudy
    visit". The extractor walks the XML in document order, looks through
    sdt wrappers at every level, and reaches a nested table exactly once
    (#921 review)."""

    NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

    def _sdt(self, inner):
        from docx.oxml import parse_xml
        return parse_xml(f'<w:sdt {self.NS}><w:sdtPr><w:tag w:val="goog_rdk_1"/></w:sdtPr>'
                         f'<w:sdtContent>{inner}</w:sdtContent></w:sdt>')

    def _docx(self, path):
        import docx
        d = docx.Document()
        d.add_paragraph("Before the table.")
        t = d.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "left"
        p = t.rows[0].cells[1].paragraphs[0]
        p.add_run("tudy visit will take hours.")
        # a run-level sdt inside a cell paragraph
        p._p.insert(list(p._p).index(p.runs[0]._r), self._sdt('<w:r><w:t xml:space="preserve">The s</w:t></w:r>'))
        # a row-level sdt wrapping a whole cell: the ☒ answer cell
        row = t.rows[1]._tr
        row.remove(row.findall(f"{{{self.NS.split(chr(34))[1]}}}tc")[1])
        row.append(self._sdt('<w:tc><w:p><w:r><w:t>☒ Yes</w:t></w:r></w:p></w:tc>'))
        t.rows[1].cells[0].text = "Multi-site study"
        # a nested table inside a cell
        inner = t.rows[0].cells[0].add_table(rows=1, cols=2)
        inner.rows[0].cells[0].text = "CES-D-10 survey"; inner.rows[0].cells[1].text = "depression"
        d.add_paragraph("After the table.")
        d.save(path)

    def test_sdt_runs_cells_and_nested_tables_are_kept_once_in_document_order(self):
        import tempfile
        from pathlib import Path
        from src.download.preprocess_sources import extract_docx_text
        with tempfile.TemporaryDirectory() as tmp:
            f = Path(tmp) / "doc.docx"
            self._docx(f)
            text = extract_docx_text(f)
        self.assertIn("The study visit will take hours.", text)
        self.assertNotIn("\ntudy", text)
        self.assertEqual(text.count("CES-D-10 survey"), 1)                 # nested table once
        self.assertIn("Multi-site study\t☒ Yes", text)                     # row-level sdt cell kept
        lines = text.splitlines()
        self.assertEqual(lines[0], "Before the table.")
        self.assertTrue(lines[1].startswith("left"))
        self.assertEqual(lines[-1], "After the table.")


if __name__ == '__main__':
    unittest.main()
