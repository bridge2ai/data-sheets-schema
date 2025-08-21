#!/usr/bin/env python3
"""
Debug extractor script for docx and xlsx files.

This script attempts to extract text content from docx and xlsx files
and save it to text files for debugging purposes.
"""

import os
import sys
from pathlib import Path
import traceback

import docx
from docx import Document
from openpyxl import load_workbook


def extract_docx(input_file, output_file):
    """Extract text content from a docx file."""
    try:
        doc = Document(input_file)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            # Document info
            f.write(f"Document: {os.path.basename(input_file)}\n")
            f.write("=" * 50 + "\n\n")
            
            # Extract all paragraphs with styles
            f.write("PARAGRAPHS:\n")
            f.write("-" * 50 + "\n")
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    style = para.style.name if para.style else "No Style"
                    f.write(f"[{i}] Style: {style}\n")
                    f.write(f"Text: {text}\n\n")
            
            # Extract all tables
            f.write("\nTABLES:\n")
            f.write("-" * 50 + "\n")
            for i, table in enumerate(doc.tables):
                f.write(f"Table {i+1}:\n")
                
                # Get table content
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    
                    f.write(f"Row {row_idx}: {' | '.join(row_text)}\n")
                f.write("\n")
                
            # Extract document properties
            f.write("\nDOCUMENT PROPERTIES:\n")
            f.write("-" * 50 + "\n")
            core_props = doc.core_properties
            try:
                f.write(f"Title: {core_props.title}\n")
                f.write(f"Author: {core_props.author}\n")
                f.write(f"Created: {core_props.created}\n")
                f.write(f"Modified: {core_props.modified}\n")
            except AttributeError:
                f.write("Could not access core properties\n")
        
        print(f"Successfully extracted content from {input_file} to {output_file}")
        return True
    except Exception as e:
        print(f"Error extracting content from {input_file}: {e}")
        traceback.print_exc()
        return False


def extract_xlsx(input_file, output_file):
    """Extract text content from an xlsx file."""
    try:
        wb = load_workbook(filename=input_file, data_only=True)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            # Document info
            f.write(f"Spreadsheet: {os.path.basename(input_file)}\n")
            f.write("=" * 50 + "\n\n")
            
            # Process each sheet
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                f.write(f"SHEET: {sheet_name}\n")
                f.write("-" * 50 + "\n")
                
                # Get sheet dimensions
                f.write(f"Dimensions: {sheet.dimensions}\n")
                f.write(f"Max Row: {sheet.max_row}, Max Column: {sheet.max_column}\n\n")
                
                # Extract data
                for row_idx in range(1, min(sheet.max_row + 1, 101)):  # Limit to 100 rows for debug
                    row_data = []
                    for col_idx in range(1, sheet.max_column + 1):
                        cell = sheet.cell(row=row_idx, column=col_idx)
                        cell_value = cell.value if cell.value is not None else ""
                        row_data.append(str(cell_value))
                    
                    f.write(f"Row {row_idx}: {' | '.join(row_data)}\n")
                
                if sheet.max_row > 100:
                    f.write("... (truncated after 100 rows)\n")
                
                f.write("\n\n")
        
        print(f"Successfully extracted content from {input_file} to {output_file}")
        return True
    except Exception as e:
        print(f"Error extracting content from {input_file}: {e}")
        traceback.print_exc()
        return False


def main():
    """Process files in GC_data_sheets directory."""
    script_dir = Path(__file__).parent
    repo_root = script_dir.parent.parent
    data_dir = repo_root / "data" / "GC_data_sheets"
    debug_dir = script_dir / "debug"
    
    if not data_dir.exists():
        print(f"Error: Data directory not found: {data_dir}")
        sys.exit(1)
    
    debug_dir.mkdir(exist_ok=True)
    
    # Process docx files
    for file_path in data_dir.glob("*.docx"):
        output_file = debug_dir / f"{file_path.stem}_debug.txt"
        extract_docx(str(file_path), str(output_file))
    
    # Process xlsx files
    for file_path in data_dir.glob("*.xlsx"):
        output_file = debug_dir / f"{file_path.stem}_debug.txt"
        extract_xlsx(str(file_path), str(output_file))


if __name__ == "__main__":
    main()