#!/usr/bin/env python3
"""
Converter for data sheets files.

This module converts docx and xlsx files containing schema data to
human-readable text-only documents with proper heading and styling.
Outputs are in txt, pdf, and html formats.
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, List, Union, Any, Optional

import docx
import pandas as pd
from fpdf import FPDF
from docx import Document
from openpyxl import load_workbook


class DataSheetConverter:
    """Convert data sheets from docx/xlsx to human-readable formats."""

    def __init__(self, input_file: str):
        """
        Initialize the converter with an input file.

        Args:
            input_file: Path to the input docx or xlsx file
        """
        self.input_file = input_file
        self.output_base = os.path.splitext(input_file)[0]
        self.data = {}
        self.file_type = os.path.splitext(input_file)[1].lower()

    def extract_data(self) -> Dict[str, Any]:
        """
        Extract data from the input file.

        Returns:
            Dict containing the extracted data structure
        """
        if self.file_type == '.docx':
            return self._extract_from_docx()
        elif self.file_type == '.xlsx':
            return self._extract_from_xlsx()
        else:
            raise ValueError(f"Unsupported file type: {self.file_type}")

    def _extract_from_docx(self) -> Dict[str, Any]:
        """
        Extract data from a docx file.

        Attempts to find and parse YAML-like schema text and structured content.
        The docx files contain YAML-formatted data with proper indentation.

        Returns:
            Dict containing structured data from the document
        """
        from .yaml_parser import extract_yaml_from_text
        
        doc = Document(self.input_file)
        data = {"title": os.path.basename(self.output_base), "sections": [], "key_values": {}}
        
        # First, concatenate all paragraphs to get the full text with indentation
        full_text = ""
        
        # Process all paragraphs to preserve indentation and structure
        for para in doc.paragraphs:
            text = para.text
            
            # Skip empty paragraphs
            if not text.strip():
                full_text += "\n"
                continue
            
            # For headings, ensure they're properly separated
            if para.style and para.style.name.startswith('Heading'):
                full_text += f"\n{text}\n"
            else:
                # Preserve indentation for normal paragraphs
                full_text += f"{text}\n"
        
        # Extract YAML data using our custom parser
        yaml_data = extract_yaml_from_text(full_text)
        if yaml_data:
            # Store top-level key-value pairs
            data["key_values"] = yaml_data
            
            # Try to identify sections in the document
            # A section might be any key with a dictionary or list value
            for key, value in yaml_data.items():
                if isinstance(value, dict):
                    section = {"title": key, "content": [], "subsections": []}
                    
                    # Process nested key-values as section content
                    for subkey, subvalue in value.items():
                        if isinstance(subvalue, dict):
                            # This is a subsection
                            subsection = {"title": subkey, "content": []}
                            
                            # Add content as formatted lines
                            for k, v in subvalue.items():
                                if isinstance(v, list):
                                    subsection["content"].append(f"{k}:")
                                    for item in v:
                                        if isinstance(item, dict):
                                            for ik, iv in item.items():
                                                subsection["content"].append(f"  - {ik}: {iv}")
                                        else:
                                            subsection["content"].append(f"  - {item}")
                                else:
                                    subsection["content"].append(f"{k}: {v}")
                            
                            section["subsections"].append(subsection)
                        elif isinstance(subvalue, list):
                            # This is a list within the section
                            section.setdefault("lists", {})[subkey] = []
                            
                            for item in subvalue:
                                if isinstance(item, dict):
                                    # Handle dictionaries in the list
                                    item_content = []
                                    for k, v in item.items():
                                        item_content.append(f"{k}: {v}")
                                    section["lists"][subkey].append(", ".join(item_content))
                                else:
                                    section["lists"][subkey].append(str(item))
                        else:
                            # Regular key-value content
                            section["content"].append(f"{subkey}: {subvalue}")
                    
                    data["sections"].append(section)
                elif isinstance(value, list):
                    # Add lists as separate sections
                    section = {"title": key, "content": [], "lists": {key: []}}
                    
                    for item in value:
                        if isinstance(item, dict):
                            # Handle dictionaries in the list
                            item_content = []
                            for k, v in item.items():
                                item_content.append(f"{k}: {v}")
                            section["lists"][key].append(", ".join(item_content))
                        else:
                            section["lists"][key].append(str(item))
                    
                    data["sections"].append(section)
        
        # Extract tables
        table_data = self._extract_tables_from_docx(doc)
        if table_data:
            data["tables"] = table_data

        return data

    def _extract_json_from_docx(self, doc: Document) -> Optional[Dict[str, Any]]:
        """
        Extract JSON data from a docx file.

        Looks for JSON-formatted blocks in the document text.

        Args:
            doc: The Document object

        Returns:
            Parsed JSON data if found, None otherwise
        """
        # Concatenate paragraphs to look for JSON blocks
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Try to find JSON blocks using regex
        json_pattern = r"\{[\s\S]*\}"
        matches = re.findall(json_pattern, full_text)
        
        for match in matches:
            try:
                return json.loads(match)
            except json.JSONDecodeError:
                continue
                
        return None

    def _extract_tables_from_docx(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract tables from a docx file.

        Args:
            doc: The Document object

        Returns:
            List of dictionaries representing tables
        """
        tables_data = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                "title": f"Table {i+1}",
                "headers": [],
                "rows": []
            }
            
            # Extract headers (first row)
            if len(table.rows) > 0:
                table_data["headers"] = [cell.text.strip() for cell in table.rows[0].cells]
                
                # Extract data rows
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data["rows"].append(row_data)
                    
            tables_data.append(table_data)
            
        return tables_data

    def _extract_from_xlsx(self) -> Dict[str, Any]:
        """
        Extract data from an xlsx file.

        The xlsx files appear to have:
        1. Key-value pairs at the top (Field-Value columns)
        2. A metadata table with headers and rows

        Returns:
            Dict containing structured data from the spreadsheet
        """
        workbook = load_workbook(filename=self.input_file, data_only=True)
        data = {
            "title": os.path.basename(self.output_base), 
            "sheets": [],
            "key_values": {},
            "tables": []
        }

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = {
                "title": sheet_name,
                "key_values": {},
                "tables": []
            }
            
            # First look for key-value pairs at the top
            # These typically have "Field" and "Value" columns or similar pattern
            field_col = None
            value_col = None
            
            # Try to find field-value structure
            if sheet.max_row > 0:
                for cell in sheet[1]:  # First row
                    if cell.value and str(cell.value).lower() == "field":
                        field_col = cell.column
                    elif cell.value and str(cell.value).lower() == "value":
                        value_col = cell.column
            
            # Process key-value pairs if found
            if field_col and value_col:
                key_values_end_row = None
                
                # Extract key-value pairs
                for row_idx in range(2, sheet.max_row + 1):
                    key_cell = sheet.cell(row=row_idx, column=field_col)
                    value_cell = sheet.cell(row=row_idx, column=value_col)
                    
                    # If a row is empty, it might indicate the end of key-values section
                    if not key_cell.value:
                        key_values_end_row = row_idx
                        break
                    
                    key = str(key_cell.value) if key_cell.value else ""
                    value = str(value_cell.value) if value_cell.value else ""
                    
                    if key:
                        sheet_data["key_values"][key] = value
                        data["key_values"][key] = value  # Store at top level too
            
            # Look for data tables - often these follow the key-value section
            # or start with a header row
            table_start_row = None
            
            # Check each row to find likely table header rows
            for row_idx in range(1, sheet.max_row + 1):
                # Skip rows we already processed as key-value pairs
                if field_col and value_col and row_idx <= (key_values_end_row or 0):
                    continue
                
                # Check if this might be a header row - has content across columns
                cells_with_content = 0
                for col_idx in range(1, sheet.max_column + 1):
                    cell = sheet.cell(row=row_idx, column=col_idx)
                    if cell.value:
                        cells_with_content += 1
                
                # If row has multiple columns with content, treat as a header row
                if cells_with_content > 2:
                    table_start_row = row_idx
                    break
            
            # Process table if we found a header row
            if table_start_row:
                headers = []
                for col_idx in range(1, sheet.max_column + 1):
                    cell = sheet.cell(row=table_start_row, column=col_idx)
                    headers.append(str(cell.value) if cell.value else "")
                
                # Remove empty headers at the end
                while headers and not headers[-1]:
                    headers.pop()
                
                if headers:  # Only process if we have valid headers
                    table = {
                        "title": f"Table in {sheet_name}",
                        "headers": headers,
                        "rows": []
                    }
                    
                    # Extract rows
                    for row_idx in range(table_start_row + 1, sheet.max_row + 1):
                        row_data = []
                        row_has_content = False
                        
                        for col_idx in range(1, len(headers) + 1):
                            cell = sheet.cell(row=row_idx, column=col_idx)
                            cell_value = str(cell.value) if cell.value is not None else ""
                            row_data.append(cell_value)
                            if cell_value:
                                row_has_content = True
                        
                        if row_has_content:
                            table["rows"].append(row_data)
                    
                    sheet_data["tables"].append(table)
                    data["tables"].append(table)  # Store at top level too
            
            data["sheets"].append(sheet_data)
        
        return data

    def generate_txt(self) -> str:
        """
        Generate a formatted text version of the data.

        Returns:
            Path to the generated text file
        """
        data = self.extract_data()
        output_file = f"{self.output_base}.txt"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            # Write title
            f.write(f"{data['title'].upper()}\n")
            f.write("=" * len(data['title']) + "\n\n")
            
            # Format and write YAML content
            if 'key_values' in data and data['key_values']:
                from .yaml_parser import convert_yaml_to_structured_text
                
                f.write("DATASET INFORMATION\n")
                f.write("-" * 19 + "\n\n")
                
                if 'Dataset' in data['key_values']:
                    # This is the common format from docx files - one top level "Dataset" key
                    yaml_content = convert_yaml_to_structured_text(data['key_values']['Dataset'], indent=0)
                    f.write(yaml_content)
                else:
                    # For other formats, write all top-level key-values
                    yaml_content = convert_yaml_to_structured_text(data['key_values'], indent=0)
                    f.write(yaml_content)
                
                f.write("\n\n")
                
            # Process content (if any is not in a section)
            if 'content' in data and data['content']:
                f.write("CONTENT\n")
                f.write("-" * 7 + "\n\n")
                
                for item in data['content']:
                    f.write(f"{item}\n")
                
                f.write("\n")
            
            # Process structured content from docx
            if 'sections' in data and data['sections']:
                for section in data['sections']:
                    # Write section heading
                    f.write(f"{section['title'].upper()}\n")
                    f.write("=" * len(section['title']) + "\n\n")
                    
                    # Write section content
                    for item in section['content']:
                        f.write(f"{item}\n")
                    
                    # Write section lists if any
                    if 'lists' in section and section['lists']:
                        f.write("\n")
                        for list_key, items in section['lists'].items():
                            f.write(f"{list_key}:\n")
                            for item in items:
                                f.write(f"  - {item}\n")
                            f.write("\n")
                    
                    # Write subsections (if any)
                    if 'subsections' in section:
                        for subsection in section['subsections']:
                            f.write(f"\n{subsection['title']}\n")
                            f.write("-" * len(subsection['title']) + "\n\n")
                            
                            for item in subsection['content']:
                                f.write(f"{item}\n")
                            
                            # Write subsection lists if any
                            if 'lists' in subsection and subsection['lists']:
                                f.write("\n")
                                for list_key, items in subsection['lists'].items():
                                    f.write(f"{list_key}:\n")
                                    for item in items:
                                        f.write(f"  - {item}\n")
                                    f.write("\n")
                    
                    f.write("\n")
            
            # Process tables
            if 'tables' in data and data['tables']:
                f.write("TABLES\n")
                f.write("=" * 6 + "\n\n")
                
                for table in data['tables']:
                    if 'title' in table:
                        f.write(f"{table['title']}\n")
                        f.write("-" * len(table['title']) + "\n")
                    
                    # Calculate column widths
                    col_widths = [len(h) for h in table['headers']]
                    for row in table['rows']:
                        for i, cell in enumerate(row):
                            if i < len(col_widths):
                                col_widths[i] = max(col_widths[i], len(str(cell)))
                    
                    # Write headers with padding
                    header_row = " | ".join(f"{h:{w}s}" for h, w in zip(table['headers'], col_widths))
                    f.write(header_row + "\n")
                    f.write("-" * len(header_row) + "\n")
                    
                    # Write rows with padding
                    for row in table['rows']:
                        padded_row = []
                        for i, cell in enumerate(row):
                            if i < len(col_widths):
                                padded_row.append(f"{str(cell):{col_widths[i]}s}")
                            else:
                                padded_row.append(str(cell))
                        f.write(" | ".join(padded_row) + "\n")
                    
                    f.write("\n")
            
            # Process sheets from xlsx (though we've already extracted most data at top level)
            if 'sheets' in data and data['sheets']:
                for sheet in data['sheets']:
                    # Only process sheet-specific data not already shown
                    if ('key_values' in sheet and sheet['key_values'] and 
                        sheet['key_values'] != data.get('key_values')):
                        
                        f.write(f"Sheet: {sheet['title']} - Metadata\n")
                        f.write("=" * (len(sheet['title']) + 18) + "\n\n")
                        
                        for key in sorted(sheet['key_values'].keys()):
                            value = sheet['key_values'][key]
                            f.write(f"{key}: {value}\n")
                        
                        f.write("\n")
                
        return output_file

    def generate_html(self) -> str:
        """
        Generate an HTML version of the data.

        Returns:
            Path to the generated HTML file
        """
        import yaml
        from .yaml_parser import format_yaml_for_html, get_yaml_css_styles, format_complex_value
        
        data = self.extract_data()
        output_file = f"{self.output_base}.html"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            # Write HTML header
            f.write("""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{0}</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 0; padding: 20px; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #3498db; margin-top: 30px; }}
        h3 {{ color: #2980b9; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ padding: 12px; text-align: left; border-bottom: 1px solid #ddd; }}
        th {{ background-color: #f2f2f2; font-weight: bold; }}
        tr:hover {{ background-color: #f5f5f5; }}
        pre {{ white-space: pre-wrap; }}
        .key-value {{ background: #f8f9fa; border: 1px solid #ddd; padding: 15px; margin: 20px 0; }}
        {1}
    </style>
</head>
<body>
    <h1>{0}</h1>
""".format(data['title'], get_yaml_css_styles()))
            
            # Format and write YAML content
            if 'key_values' in data and data['key_values']:
                f.write("<h2>Dataset Information</h2>\n")
                f.write("<div class=\"yaml-block\">\n")
                
                if 'Dataset' in data['key_values']:
                    # This is the common format from docx files - one top level "Dataset" key
                    yaml_html = format_yaml_for_html(data['key_values']['Dataset'])
                    f.write(yaml_html)
                else:
                    # For other formats, write all top-level key-values
                    yaml_content = {}
                    for key, value in data['key_values'].items():
                        # Try to parse complex values (e.g., dictionaries in string form)
                        if isinstance(value, str) and value.startswith('{') and value.endswith('}'):
                            try:
                                # Convert string repr of dict to actual dict
                                complex_value = eval(value)
                                if isinstance(complex_value, dict):
                                    yaml_content[key] = complex_value
                                    continue
                            except:
                                pass
                        yaml_content[key] = value
                    
                    yaml_html = format_yaml_for_html(yaml_content)
                    f.write(yaml_html)
                
                f.write("</div>\n")
                
            # Process content (if any is not in a section)
            if 'content' in data and data['content']:
                f.write("<h2>Content</h2>\n")
                f.write("<div class=\"yaml-block\">\n")
                
                for item in data['content']:
                    # Try to handle key-value formats
                    if ': ' in item:
                        key, value = item.split(': ', 1)
                        # Try to parse complex values
                        try:
                            if value.startswith('{') and value.endswith('}'):
                                complex_value = eval(value)
                                if isinstance(complex_value, dict):
                                    f.write(f'<div class="indent-0"><span class="key">{key}</span>:</div>\n')
                                    f.write(format_yaml_for_html(complex_value, indent=1))
                                    continue
                        except:
                            pass
                        f.write(f'<div><span class="key">{key}</span>: <span class="value">{value}</span></div>\n')
                    else:
                        f.write(f"<div>{item}</div>\n")
                
                f.write("</div>\n")
            
            # Process structured content from docx
            if 'sections' in data and data['sections']:
                for section in data['sections']:
                    # Write section heading
                    f.write(f"<h2>{section['title']}</h2>\n")
                    
                    # Format section content
                    if section['content']:
                        f.write("<div class=\"yaml-block\">\n")
                        
                        for item in section['content']:
                            # Check if item is a key-value pair
                            if ': ' in item:
                                key, value = item.split(': ', 1)
                                # Try to parse complex values
                                try:
                                    if value.startswith('{') and value.endswith('}'):
                                        complex_value = eval(value)
                                        if isinstance(complex_value, dict):
                                            f.write(f'<div class="indent-0"><span class="key">{key}</span>:</div>\n')
                                            f.write(format_yaml_for_html(complex_value, indent=1))
                                            continue
                                except:
                                    pass
                                f.write(f'<div class="indent-0"><span class="key">{key}</span>: <span class="value">{value}</span></div>\n')
                            else:
                                f.write(f'<div>{item}</div>\n')
                        
                        f.write("</div>\n")
                    
                    # Lists (e.g., keywords, etc.)
                    if 'lists' in section and section['lists']:
                        f.write("<div class=\"yaml-block\">\n")
                        
                        for list_key, items in section['lists'].items():
                            f.write(f'<div class="indent-0"><span class="key">{list_key}</span>:</div>\n')
                            
                            for item in items:
                                # Try to parse complex items
                                try:
                                    if isinstance(item, str) and item.startswith('{') and item.endswith('}'):
                                        complex_value = eval(item)
                                        if isinstance(complex_value, dict):
                                            f.write(f'<div class="indent-1">&nbsp;&nbsp;- </div>\n')
                                            f.write(format_yaml_for_html(complex_value, indent=2))
                                            continue
                                except:
                                    pass
                                f.write(f'<div class="indent-1">&nbsp;&nbsp;- <span class="value">{item}</span></div>\n')
                        
                        f.write("</div>\n")
                    
                    # Write subsections
                    if 'subsections' in section:
                        for subsection in section['subsections']:
                            f.write(f"<h3>{subsection['title']}</h3>\n")
                            
                            # Format subsection content
                            if subsection['content']:
                                f.write("<div class=\"yaml-block\">\n")
                                
                                for item in subsection['content']:
                                    # Check if item is a key-value pair
                                    if ': ' in item:
                                        key, value = item.split(': ', 1)
                                        # Try to parse complex values
                                        try:
                                            if value.startswith('{') and value.endswith('}'):
                                                complex_value = eval(value)
                                                if isinstance(complex_value, dict):
                                                    f.write(f'<div class="indent-0"><span class="key">{key}</span>:</div>\n')
                                                    f.write(format_yaml_for_html(complex_value, indent=1))
                                                    continue
                                        except:
                                            pass
                                        f.write(f'<div class="indent-0"><span class="key">{key}</span>: <span class="value">{value}</span></div>\n')
                                    else:
                                        f.write(f'<div>{item}</div>\n')
                                
                                f.write("</div>\n")
                            
                            # Lists (e.g., keywords, etc.)
                            if 'lists' in subsection and subsection['lists']:
                                f.write("<div class=\"yaml-block\">\n")
                                
                                for list_key, items in subsection['lists'].items():
                                    f.write(f'<div class="indent-0"><span class="key">{list_key}</span>:</div>\n')
                                    
                                    for item in items:
                                        f.write(f'<div class="indent-1">&nbsp;&nbsp;- <span class="value">{item}</span></div>\n')
                                
                                f.write("</div>\n")
            
            # Process tables
            if 'tables' in data and data['tables']:
                f.write("<h2>Tables</h2>\n")
                
                for table in data['tables']:
                    if 'title' in table:
                        f.write(f"<h3>{table['title']}</h3>\n")
                    
                    f.write("<table>\n  <tr>\n")
                    
                    # Write headers
                    for header in table['headers']:
                        f.write(f"    <th>{header}</th>\n")
                    f.write("  </tr>\n")
                    
                    # Write rows
                    for row in table['rows']:
                        f.write("  <tr>\n")
                        for i, cell in enumerate(row):
                            # Handle case where row might have fewer cells than headers
                            if i < len(table['headers']):
                                # Try to parse complex values
                                try:
                                    if isinstance(cell, str) and cell.startswith('{') and cell.endswith('}'):
                                        complex_value = eval(cell)
                                        if isinstance(complex_value, dict):
                                            yaml_html = format_yaml_for_html({table['headers'][i]: complex_value})
                                            f.write(f"    <td><div class=\"yaml-block\">{yaml_html}</div></td>\n")
                                            continue
                                except:
                                    pass
                                f.write(f"    <td>{cell}</td>\n")
                        f.write("  </tr>\n")
                    
                    f.write("</table>\n")
            
            # Process sheets from xlsx (though we've already extracted most data at top level)
            if 'sheets' in data and data['sheets']:
                for sheet in data['sheets']:
                    # Only process sheet-specific data not already shown
                    if ('key_values' in sheet and sheet['key_values'] and 
                        sheet['key_values'] != data.get('key_values')):
                        
                        f.write(f"<h3>Sheet: {sheet['title']} - Metadata</h3>\n")
                        f.write("<div class=\"yaml-block\">\n")
                        
                        yaml_content = {}
                        for key, value in sheet['key_values'].items():
                            # Try to parse complex values (e.g., dictionaries in string form)
                            if isinstance(value, str) and value.startswith('{') and value.endswith('}'):
                                try:
                                    complex_value = eval(value)
                                    if isinstance(complex_value, dict):
                                        yaml_content[key] = complex_value
                                        continue
                                except:
                                    pass
                            yaml_content[key] = value
                        
                        yaml_html = format_yaml_for_html(yaml_content)
                        f.write(yaml_html)
                        
                        f.write("</div>\n")
                
            # Write HTML footer
            f.write("</body>\n</html>")
                
        return output_file

    def generate_pdf(self) -> str:
        """
        Generate a PDF version of the data.

        Returns:
            Path to the generated PDF file
        """
        import yaml
        from .yaml_parser import format_yaml_for_pdf, format_complex_value
        
        data = self.extract_data()
        output_file = f"{self.output_base}.pdf"
        
        pdf = FPDF()
        pdf.add_page()
        
        # Set up fonts
        pdf.set_font("Arial", "B", 16)
        
        # Title
        pdf.cell(0, 10, data['title'], 0, 1, 'C')
        pdf.ln(5)
        
        # Format and write YAML content
        if 'key_values' in data and data['key_values']:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "DATASET INFORMATION", 0, 1)
            pdf.ln(2)
            
            y_pos = pdf.get_y()
            
            if 'Dataset' in data['key_values']:
                # This is the common format from docx files - one top level "Dataset" key
                y_pos = format_yaml_for_pdf(data['key_values']['Dataset'], pdf, 10, y_pos)
            else:
                # Process complex values
                yaml_content = {}
                for key, value in data['key_values'].items():
                    # Try to parse complex values (e.g., dictionaries in string form)
                    if isinstance(value, str) and value.startswith('{') and value.endswith('}'):
                        try:
                            # Convert string repr of dict to actual dict
                            complex_value = eval(value)
                            if isinstance(complex_value, dict):
                                yaml_content[key] = complex_value
                                continue
                        except:
                            pass
                    yaml_content[key] = value
                
                y_pos = format_yaml_for_pdf(yaml_content, pdf, 10, y_pos)
            
            pdf.set_y(y_pos + 5)
        
        # Process content (if any is not in a section)
        if 'content' in data and data['content']:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "CONTENT", 0, 1)
            
            y_pos = pdf.get_y()
            
            content_data = {}
            for item in data['content']:
                # Try to handle key-value formats
                if ': ' in item:
                    key, value = item.split(': ', 1)
                    # Try to parse complex values
                    try:
                        if value.startswith('{') and value.endswith('}'):
                            complex_value = eval(value)
                            if isinstance(complex_value, dict):
                                content_data[key] = complex_value
                                continue
                    except:
                        pass
                    content_data[key] = value
                else:
                    # Just add as plain text
                    pdf.set_font("Arial", "", 10)
                    pdf.multi_cell(0, 10, item)
            
            if content_data:
                y_pos = format_yaml_for_pdf(content_data, pdf, 10, y_pos)
            
            pdf.set_y(y_pos + 5)
        
        # Process structured content from docx
        if 'sections' in data and data['sections']:
            for section in data['sections']:
                # Section heading
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, section['title'].upper(), 0, 1)
                
                y_pos = pdf.get_y()
                
                # Format section content
                section_data = {}
                
                # Process regular content items
                for item in section['content']:
                    # Parse key-value pairs
                    if ': ' in item:
                        key, value = item.split(': ', 1)
                        # Try to parse complex values
                        try:
                            if value.startswith('{') and value.endswith('}'):
                                complex_value = eval(value)
                                if isinstance(complex_value, dict):
                                    section_data[key] = complex_value
                                    continue
                        except:
                            pass
                        section_data[key] = value
                    else:
                        # Plain text, just add directly
                        pdf.set_font("Arial", "", 10)
                        pdf.multi_cell(0, 10, item)
                
                # Format as YAML if we have structured data
                if section_data:
                    y_pos = format_yaml_for_pdf(section_data, pdf, 10, y_pos)
                    pdf.set_y(y_pos)
                
                # Process lists
                if 'lists' in section and section['lists']:
                    for list_key, items in section['lists'].items():
                        pdf.set_font("Arial", "B", 10)
                        pdf.cell(0, 10, f"{list_key}:", 0, 1)
                        
                        # Handle each item
                        list_items = []
                        for item in items:
                            # Try to parse complex items
                            try:
                                if isinstance(item, str) and item.startswith('{') and item.endswith('}'):
                                    complex_value = eval(item)
                                    if isinstance(complex_value, dict):
                                        y_pos = pdf.get_y()
                                        pdf.cell(10, 10, "•", 0, 0)
                                        y_pos = format_yaml_for_pdf(complex_value, pdf, 20, y_pos)
                                        pdf.set_y(y_pos)
                                        continue
                            except:
                                pass
                            
                            # Regular list item
                            pdf.set_font("Arial", "", 10)
                            pdf.cell(10, 10, "•", 0, 0)
                            pdf.multi_cell(0, 10, item)
                
                # Process subsections
                if 'subsections' in section:
                    for subsection in section['subsections']:
                        pdf.set_font("Arial", "B", 12)
                        pdf.cell(0, 10, subsection['title'], 0, 1)
                        
                        y_pos = pdf.get_y()
                        
                        # Format subsection content
                        subsection_data = {}
                        
                        for item in subsection['content']:
                            # Parse key-value pairs
                            if ': ' in item:
                                key, value = item.split(': ', 1)
                                # Try to parse complex values
                                try:
                                    if value.startswith('{') and value.endswith('}'):
                                        complex_value = eval(value)
                                        if isinstance(complex_value, dict):
                                            subsection_data[key] = complex_value
                                            continue
                                except:
                                    pass
                                subsection_data[key] = value
                            else:
                                # Plain text, just add directly
                                pdf.set_font("Arial", "", 10)
                                pdf.multi_cell(0, 10, item)
                        
                        # Format as YAML if we have structured data
                        if subsection_data:
                            y_pos = format_yaml_for_pdf(subsection_data, pdf, 10, y_pos)
                            pdf.set_y(y_pos)
                        
                        # Process lists
                        if 'lists' in subsection and subsection['lists']:
                            for list_key, items in subsection['lists'].items():
                                pdf.set_font("Arial", "B", 10)
                                pdf.cell(0, 10, f"{list_key}:", 0, 1)
                                
                                pdf.set_font("Arial", "", 10)
                                for item in items:
                                    pdf.cell(10, 10, "•", 0, 0)
                                    pdf.multi_cell(0, 10, item)
                
                pdf.ln(5)
        
        # Tables - simplified representation for PDF
        if 'tables' in data and data['tables']:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "TABLES", 0, 1)
            
            for table in data['tables']:
                if 'title' in table:
                    pdf.set_font("Arial", "B", 12)
                    pdf.cell(0, 10, table['title'], 0, 1)
                
                # Headers
                pdf.set_font("Arial", "B", 10)
                header_text = " | ".join(str(h) for h in table['headers'])
                if len(header_text) > 100:  # Truncate overly long headers
                    header_text = header_text[:97] + "..."
                pdf.multi_cell(0, 10, header_text)
                
                # Rows
                pdf.set_font("Arial", "", 10)
                for row in table['rows']:
                    # Check each cell for complex values
                    processed_cells = []
                    for i, cell in enumerate(row):
                        # Try to parse complex values
                        try:
                            if isinstance(cell, str) and cell.startswith('{') and cell.endswith('}'):
                                complex_value = eval(cell)
                                if isinstance(complex_value, dict):
                                    # Just use a placeholder and process complex content separately
                                    processed_cells.append("[Complex data]")
                                    
                                    # Add the complex data on a new line
                                    y_pos = pdf.get_y() + 10
                                    header = table['headers'][i] if i < len(table['headers']) else f"Column {i+1}"
                                    pdf.set_y(y_pos)
                                    pdf.set_font("Arial", "B", 8)
                                    pdf.cell(0, 10, f"{header}:", 0, 1)
                                    y_pos = format_yaml_for_pdf(complex_value, pdf, 20, y_pos + 10)
                                    pdf.set_y(y_pos)
                                    continue
                        except:
                            pass
                        processed_cells.append(str(cell))
                    
                    # Format the row
                    row_text = " | ".join(processed_cells[:len(table['headers'])])
                    if len(row_text) > 100:  # Truncate overly long rows
                        row_text = row_text[:97] + "..."
                    pdf.multi_cell(0, 10, row_text)
                
                pdf.ln(5)
        
        # Process sheets from xlsx (though we've already extracted most data at top level)
        if 'sheets' in data and data['sheets']:
            for sheet in data['sheets']:
                # Only process sheet-specific data not already shown
                if ('key_values' in sheet and sheet['key_values'] and 
                    sheet['key_values'] != data.get('key_values')):
                    
                    pdf.set_font("Arial", "B", 12)
                    pdf.cell(0, 10, f"Sheet: {sheet['title']} - Metadata", 0, 1)
                    
                    y_pos = pdf.get_y()
                    
                    # Process complex values
                    yaml_content = {}
                    for key, value in sheet['key_values'].items():
                        # Try to parse complex values
                        try:
                            if isinstance(value, str) and value.startswith('{') and value.endswith('}'):
                                complex_value = eval(value)
                                if isinstance(complex_value, dict):
                                    yaml_content[key] = complex_value
                                    continue
                        except:
                            pass
                        yaml_content[key] = value
                    
                    y_pos = format_yaml_for_pdf(yaml_content, pdf, 10, y_pos)
                    pdf.set_y(y_pos + 5)
        
        # Save the PDF
        pdf.output(output_file)
        return output_file

    def convert(self) -> List[str]:
        """
        Convert the input file to all output formats.

        Returns:
            List of paths to the generated files
        """
        txt_file = self.generate_txt()
        html_file = self.generate_html()
        pdf_file = self.generate_pdf()
        
        return [txt_file, html_file, pdf_file]


def process_directory(directory: str) -> None:
    """
    Process all docx and xlsx files in a directory.

    Args:
        directory: Path to the directory containing files to process
    """
    input_dir = Path(directory)
    
    # Create output directories if they don't exist
    txt_dir = input_dir / "txt"
    html_dir = input_dir / "html"
    pdf_dir = input_dir / "pdf"
    
    txt_dir.mkdir(exist_ok=True)
    html_dir.mkdir(exist_ok=True)
    pdf_dir.mkdir(exist_ok=True)
    
    # Process each file
    for file_path in input_dir.glob("*.docx"):
        try:
            print(f"Processing {file_path.name}...")
            converter = DataSheetConverter(str(file_path))
            output_files = converter.convert()
            print(f"Generated: {', '.join(output_files)}")
        except Exception as e:
            print(f"Error processing {file_path.name}: {e}")
    
    for file_path in input_dir.glob("*.xlsx"):
        try:
            print(f"Processing {file_path.name}...")
            converter = DataSheetConverter(str(file_path))
            output_files = converter.convert()
            print(f"Generated: {', '.join(output_files)}")
        except Exception as e:
            print(f"Error processing {file_path.name}: {e}")


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        input_path = sys.argv[1]
        if os.path.isdir(input_path):
            process_directory(input_path)
        elif os.path.isfile(input_path):
            file_ext = os.path.splitext(input_path)[1].lower()
            if file_ext in ('.docx', '.xlsx'):
                try:
                    converter = DataSheetConverter(input_path)
                    output_files = converter.convert()
                    print(f"Generated: {', '.join(output_files)}")
                except Exception as e:
                    print(f"Error processing {input_path}: {e}")
            else:
                print(f"Unsupported file type: {file_ext}")
        else:
            print(f"Path not found: {input_path}")
    else:
        print("Usage: python converter.py <input_file_or_directory>")
        sys.exit(1)