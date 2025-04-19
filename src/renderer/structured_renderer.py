#!/usr/bin/env python3
"""
Structured renderer for Datasheets.

This module provides rendering with proper Collection-based organization
of the content following the D4D Collections structure.
"""

import os
import sys
import io
import re
import yaml
import json
from pathlib import Path
import logging
from collections import OrderedDict
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import html

# Configure stdout/stderr to use UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# Fix environment encoding
os.environ['PYTHONIOENCODING'] = 'utf-8'

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Define all collection names in their proper order
COLLECTION_ORDER = [
    "D4D-Motivation",
    "D4D-Collection",
    "D4D-Composition",
    "D4D-Preprocessing",
    "D4D-Uses",
    "D4D-Distribution",
    "D4D-Maintenance",
    "D4D-Ethics-and-Governance"
]

# Define collection titles
COLLECTION_TITLES = {
    "D4D-Motivation": "Motivation",
    "D4D-Collection": "Collection Process",
    "D4D-Composition": "Composition",
    "D4D-Preprocessing": "Preprocessing, Cleaning, and Labeling",
    "D4D-Uses": "Uses",
    "D4D-Distribution": "Distribution",
    "D4D-Maintenance": "Maintenance",
    "D4D-Ethics-and-Governance": "Ethics and Governance"
}

# Define collection descriptions
COLLECTION_DESCRIPTIONS = {
    "D4D-Motivation": "Questions about the purpose and creation of the dataset",
    "D4D-Collection": "Questions about how the data was collected",
    "D4D-Composition": "Questions about what the dataset contains",
    "D4D-Preprocessing": "Questions about data preparation",
    "D4D-Uses": "Questions about dataset uses",
    "D4D-Distribution": "Questions about dataset distribution",
    "D4D-Maintenance": "Questions about dataset maintenance",
    "D4D-Ethics-and-Governance": "Questions about ethical review and governance of the dataset"
}

def sanitize_text(text):
    """Replace problematic Unicode characters with ASCII equivalents."""
    if not text:
        return text
        
    # Define replacements for problematic characters
    replacements = {
        '\u2022': '*',    # bullet
        '\u2019': "'",    # right single quotation mark
        '\u2018': "'",    # left single quotation mark
        '\u201C': '"',    # left double quotation mark
        '\u201D': '"',    # right double quotation mark
        '\u2013': '-',    # en dash
        '\u2014': '--',   # em dash
        '\u2026': '...',  # horizontal ellipsis
        '\u00A0': ' ',    # non-breaking space
    }
    
    # Apply replacements
    for unicode_char, ascii_char in replacements.items():
        text = text.replace(unicode_char, ascii_char)
        
    return text

def extract_and_structure_docx(docx_path):
    """
    Extract text from DOCX and structure it by collection.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        dict: Data structured by collection
    """
    try:
        doc = docx.Document(docx_path)
        all_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text.append(sanitize_text(text))
                
        full_text = "\n".join(all_text)
        
        # Try to identify which collection this document belongs to
        collection_name = None
        docx_path_str = str(docx_path).lower()
        for collection in COLLECTION_ORDER:
            short_name = collection.split('-')[-1].lower()
            title = COLLECTION_TITLES[collection].lower()
            
            if short_name in docx_path_str or title in docx_path_str:
                collection_name = collection
                break
        
        # If we couldn't identify the collection, make a guess
        if not collection_name:
            # Look through the text for clues
            for collection in COLLECTION_ORDER:
                title = COLLECTION_TITLES[collection].lower()
                if title in full_text.lower():
                    collection_name = collection
                    break
            
            # If still not found, default to first collection
            if not collection_name:
                collection_name = COLLECTION_ORDER[0]
        
        # Now parse the content looking for questions and answers
        result = {
            "collection": collection_name,
            "title": COLLECTION_TITLES.get(collection_name, "Dataset Information"),
            "description": COLLECTION_DESCRIPTIONS.get(collection_name, ""),
            "questions": OrderedDict()
        }
        
        # Split into paragraphs and try to identify questions and answers
        paragraphs = full_text.split('\n')
        current_question = None
        current_answer = []
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            # Check if this looks like a question
            if para.endswith('?') or para.startswith('Q:') or re.match(r'^\d+[\.\)]', para):
                # If we were collecting an answer, save it
                if current_question and current_answer:
                    result["questions"][current_question] = "\n".join(current_answer)
                    current_answer = []
                
                # This is a new question
                current_question = para
            elif current_question:
                # This is part of an answer
                current_answer.append(para)
        
        # Add the last question/answer if any
        if current_question and current_answer:
            result["questions"][current_question] = "\n".join(current_answer)
            
        return result
            
    except Exception as e:
        logger.error(f"Error extracting and structuring {docx_path}: {str(e)}")
        return None

def merge_structured_data(file_data_list):
    """
    Merge structured data from multiple files.
    
    Args:
        file_data_list: List of structured data from files
        
    Returns:
        dict: Merged data organized by collection
    """
    result = OrderedDict()
    
    # Initialize collections in the correct order
    for collection in COLLECTION_ORDER:
        result[collection] = {
            "title": COLLECTION_TITLES.get(collection, collection),
            "description": COLLECTION_DESCRIPTIONS.get(collection, ""),
            "questions": OrderedDict()
        }
    
    # Go through each file's data and add questions to appropriate collection
    for file_data in file_data_list:
        if not file_data or "collection" not in file_data:
            continue
            
        collection = file_data["collection"]
        if collection not in result:
            # Should not happen with our predefined collections, but just in case
            result[collection] = {
                "title": file_data.get("title", collection),
                "description": file_data.get("description", ""),
                "questions": OrderedDict()
            }
            
        # Add questions
        for question, answer in file_data.get("questions", {}).items():
            result[collection]["questions"][question] = answer
    
    return result

def render_as_markdown(data, output_path=None):
    """
    Render structured data as Markdown.
    
    Args:
        data: Structured data organized by collection
        output_path: Output file path (optional)
        
    Returns:
        str: Markdown text
    """
    lines = ["# Datasheets for Datasets\n"]
    
    # Add each collection
    for collection_id, collection in data.items():
        if not collection.get("questions"):
            continue  # Skip empty collections
            
        lines.append(f"## {collection['title']}")
        lines.append(collection["description"])
        lines.append("")
        
        # Add each question and answer
        for question, answer in collection["questions"].items():
            lines.append(f"### {question}")
            lines.append(answer)
            lines.append("")
            
        lines.append("")
    
    # Join lines into text
    markdown_text = "\n".join(lines)
    
    # Save to file if path provided
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_text)
            
    return markdown_text

def render_as_html(data, output_path=None):
    """
    Render structured data as HTML.
    
    Args:
        data: Structured data organized by collection
        output_path: Output file path (optional)
        
    Returns:
        str: HTML text
    """
    lines = ["""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Datasheets for Datasets</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
        }
        h1 {
            color: #2c3e50;
            text-align: center;
            margin-bottom: 30px;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }
        h2 {
            color: #2980b9;
            margin-top: 30px;
            padding: 10px;
            background-color: #f5f5f5;
            border-left: 5px solid #3498db;
        }
        h3 {
            color: #3498db;
            margin-top: 20px;
            border-bottom: 1px solid #ddd;
        }
        .collection {
            margin-bottom: 30px;
            padding: 20px;
            background-color: #f9f9f9;
            border-radius: 5px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }
        .collection-description {
            font-style: italic;
            color: #666;
            margin-bottom: 20px;
        }
        .question {
            margin-bottom: 20px;
            padding: 15px;
            background-color: #fff;
            border-radius: 5px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }
        .question-text {
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 10px;
        }
        .answer {
            white-space: pre-line;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        th {
            background-color: #3498db;
            color: white;
            text-align: left;
            padding: 10px;
        }
        td {
            border: 1px solid #ddd;
            padding: 10px;
            vertical-align: top;
        }
        tr:nth-child(even) {
            background-color: #f2f2f2;
        }
    </style>
</head>
<body>
    <h1>Datasheets for Datasets</h1>
"""]
    
    # Add table of contents
    lines.append('    <div class="collection">')
    lines.append('        <h2>Table of Contents</h2>')
    lines.append('        <ol>')
    
    for collection_id, collection in data.items():
        if not collection.get("questions"):
            continue  # Skip empty collections
        lines.append(f'            <li><a href="#{collection_id}">{collection["title"]}</a></li>')
    
    lines.append('        </ol>')
    lines.append('    </div>')
    
    # Add each collection
    for collection_id, collection in data.items():
        if not collection.get("questions"):
            continue  # Skip empty collections
            
        lines.append(f'    <div class="collection" id="{collection_id}">')
        lines.append(f'        <h2>{collection["title"]}</h2>')
        lines.append(f'        <div class="collection-description">{collection["description"]}</div>')
        
        # Add table of questions for this collection
        lines.append('        <table>')
        lines.append('            <tr><th>Question</th><th>Answer</th></tr>')
        
        # Add each question and answer as table rows
        for question, answer in collection["questions"].items():
            # Escape HTML special characters
            question_html = html.escape(question)
            answer_html = html.escape(answer).replace('\n', '<br>')
            
            lines.append('            <tr>')
            lines.append(f'                <td class="question-text">{question_html}</td>')
            lines.append(f'                <td class="answer">{answer_html}</td>')
            lines.append('            </tr>')
        
        lines.append('        </table>')
        lines.append('    </div>')
    
    lines.append("</body>")
    lines.append("</html>")
    
    # Join lines into text
    html_text = "\n".join(lines)
    
    # Save to file if path provided
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_text)
            
    return html_text

def render_as_docx(data, output_path=None):
    """
    Render structured data as DOCX.
    
    Args:
        data: Structured data organized by collection
        output_path: Output file path (optional)
        
    Returns:
        docx.Document: DOCX document
    """
    # Create document
    doc = docx.Document()
    
    # Set styles
    style = doc.styles['Title']
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = RGBColor(44, 62, 80)  # #2c3e50
    
    style = doc.styles['Heading 1']
    style.font.size = Pt(14)
    style.font.bold = True
    style.font.color.rgb = RGBColor(41, 128, 185)  # #2980b9
    
    style = doc.styles['Heading 2']
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.color.rgb = RGBColor(52, 152, 219)  # #3498db
    
    # Add title
    doc.add_paragraph('Datasheets for Datasets', style='Title')
    
    # Add each collection
    for collection_id, collection in data.items():
        if not collection.get("questions"):
            continue  # Skip empty collections
            
        # Add collection heading
        doc.add_paragraph(collection["title"], style='Heading 1')
        
        # Add collection description
        doc.add_paragraph(collection["description"])
        
        # Add questions as a table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Set headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Question"
        header_cells[1].text = "Answer"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add each question and answer
        for question, answer in collection["questions"].items():
            row_cells = table.add_row().cells
            row_cells[0].text = question
            row_cells[1].text = answer
        
        # Add spacing
        doc.add_paragraph()
    
    # Save to file if path provided
    if output_path:
        doc.save(output_path)
            
    return doc

def process_directory(input_dir, output_dir=None):
    """
    Process all DOCX files in a directory and render structured output.
    
    Args:
        input_dir: Directory containing DOCX files
        output_dir: Output directory (optional)
        
    Returns:
        dict: Paths to output files
    """
    input_dir = Path(input_dir)
    if not input_dir.exists() or not input_dir.is_dir():
        logger.error(f"Input directory not found: {input_dir}")
        return {}
        
    # Set output directory
    if not output_dir:
        output_dir = input_dir / "output"
    else:
        output_dir = Path(output_dir)
        
    # Create output directory
    output_dir.mkdir(exist_ok=True)
    
    # Find DOCX files
    docx_files = list(input_dir.glob("*.docx"))
    if not docx_files:
        logger.error(f"No DOCX files found in {input_dir}")
        return {}
        
    logger.info(f"Found {len(docx_files)} DOCX files in {input_dir}")
    
    # Process each file and collect structured data
    file_data_list = []
    for docx_file in docx_files:
        logger.info(f"Extracting data from {docx_file.name}...")
        data = extract_and_structure_docx(docx_file)
        if data:
            file_data_list.append(data)
    
    if not file_data_list:
        logger.error("No data could be extracted from files")
        return {}
    
    # Merge data from all files
    logger.info("Merging data from all files...")
    merged_data = merge_structured_data(file_data_list)
    
    # Render outputs
    output_files = {}
    
    # Markdown
    logger.info("Rendering Markdown...")
    md_path = output_dir / "datasheet.md"
    render_as_markdown(merged_data, md_path)
    output_files["markdown"] = str(md_path)
    
    # HTML
    logger.info("Rendering HTML...")
    html_path = output_dir / "datasheet.html"
    render_as_html(merged_data, html_path)
    output_files["html"] = str(html_path)
    
    # DOCX
    logger.info("Rendering DOCX...")
    docx_path = output_dir / "datasheet.docx"
    render_as_docx(merged_data, docx_path)
    output_files["docx"] = str(docx_path)
    
    return output_files

def main():
    """Main entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Process DOCX files and render structured output"
    )
    parser.add_argument(
        "input",
        help="Input directory containing DOCX files"
    )
    parser.add_argument(
        "--output", "-o",
        help="Output directory for rendered files"
    )
    
    args = parser.parse_args()
    
    # Process files
    logger.info(f"Processing files from {args.input}...")
    output_files = process_directory(args.input, args.output)
    
    if output_files:
        logger.info("Processing complete. Output files:")
        for format_name, file_path in output_files.items():
            logger.info(f"  {format_name.upper()}: {file_path}")
        return 0
    else:
        logger.error("Processing failed.")
        return 1
    
if __name__ == "__main__":
    # When running directly, ensure we have sample data to demonstrate the outputs
    if len(sys.argv) < 2:
        print("No input directory specified. Using sample data for demonstration.")
        
        # Create sample data structure for demonstration
        sample_data = OrderedDict()
        for collection in COLLECTION_ORDER:
            sample_data[collection] = {
                "title": COLLECTION_TITLES.get(collection, collection),
                "description": COLLECTION_DESCRIPTIONS.get(collection, ""),
                "questions": OrderedDict()
            }
            
            # Add sample questions for each collection
            sample_data[collection]["questions"]["What is the purpose of this collection?"] = f"This collection covers {COLLECTION_TITLES.get(collection, collection)} information about the dataset."
            sample_data[collection]["questions"]["Why is this information important?"] = "It helps users understand the dataset better and make informed decisions about its use."
        
        # Set up an output directory
        output_dir = Path("./output")
        output_dir.mkdir(exist_ok=True)
        
        # Render all formats
        print("Rendering example outputs with sample data...")
        
        # Markdown
        md_path = output_dir / "datasheet_example.md"
        render_as_markdown(sample_data, md_path)
        print(f"  Markdown: {md_path}")
        
        # HTML
        html_path = output_dir / "datasheet_example.html"
        render_as_html(sample_data, html_path)
        print(f"  HTML: {html_path}")
        
        # DOCX
        docx_path = output_dir / "datasheet_example.docx"
        render_as_docx(sample_data, docx_path)
        print(f"  DOCX: {docx_path}")
        
        print("\nSample outputs generated successfully. Run with an input directory to process real data.")
    else:
        # Normal execution with command-line arguments
        sys.exit(main())